"""Create an editable, client-facing Word brief from one website audit."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "outputs"


def _hex_color(value: str, fallback: str) -> RGBColor:
    raw = (value or fallback).lstrip("#")
    if len(raw) != 6:
        raw = fallback.lstrip("#")
    return RGBColor.from_string(raw.upper())


def _set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color.lstrip("#").upper())


def _set_cell_margins(cell, top: int = 180, start: int = 220, bottom: int = 180, end: int = 220) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_font(run, name: str, size: float, color: RGBColor, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def _add_text(
    doc: Document,
    text: str,
    *,
    font: str,
    size: float,
    color: RGBColor,
    bold: bool = False,
    italic: bool = False,
    before: float = 0,
    after: float = 7,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    _set_font(run, font, size, color, bold=bold, italic=italic)


def _add_heading(doc: Document, text: str, display_font: str, ink: RGBColor, *, size: float = 25) -> None:
    _add_text(doc, text, font=display_font, size=size, color=ink, after=12)


def _add_bullet(doc: Document, text: str, body_font: str, ink: RGBColor) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.16)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    _set_font(run, body_font, 11, ink)


def _add_page_title(doc: Document, kicker: str, title: str, display_font: str, body_font: str, ink: RGBColor, accent: RGBColor) -> None:
    _add_text(doc, kicker.upper(), font=body_font, size=8.5, color=accent, bold=True, after=8)
    _add_heading(doc, title, display_font, ink, size=26)


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def render_editable_docx(audit: dict, report_id: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{report_id}.docx"
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.68)
    section.right_margin = Inches(0.78)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    display_font = audit["brand"].get("display_font") or "Georgia"
    body_font = audit["brand"].get("body_font") or "Arial"
    ink = _hex_color("#1a1a1a", "#1a1a1a")
    muted = _hex_color("#625f5b", "#625f5b")
    white = _hex_color("#ffffff", "#ffffff")
    brand = _hex_color(audit["brand"].get("primary_hex"), "#1a1a1a")
    accent = _hex_color(audit["brand"].get("accent_hex"), "#b89a6a")

    normal = doc.styles["Normal"]
    normal.font.name = body_font
    normal._element.rPr.rFonts.set(qn("w:ascii"), body_font)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), body_font)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_font(header.add_run(f"{audit['brand']['name']} | Website review"), body_font, 8.5, muted, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(footer.add_run("Editable working document"), body_font, 8, muted)

    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.autofit = False
    cover.columns[0].width = Inches(9.4)
    cell = cover.cell(0, 0)
    cell.width = Inches(9.4)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_fill(cell, audit["brand"].get("primary_hex") or "#1a1a1a")
    _set_cell_margins(cell, top=700, start=520, bottom=700, end=520)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(12)
    _set_font(paragraph.add_run("WEBSITE REVIEW"), body_font, 9, accent, bold=True)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(14)
    _set_font(paragraph.add_run(audit["brand"]["name"]), display_font, 36, white)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    _set_font(paragraph.add_run(audit["client_review"]["headline"]), body_font, 15, white)
    paragraph = cell.add_paragraph()
    _set_font(
        paragraph.add_run(f"{audit['url']} | {audit['generated_at'][:10]}"),
        body_font,
        8.5,
        white,
    )

    _page_break(doc)
    _add_page_title(doc, "The short answer", audit["client_review"]["headline"], display_font, body_font, ink, accent)
    _add_text(
        doc,
        f"Current conversion readiness: {audit['conversion']['score']}/100. The first issue to solve is {audit['conversion']['root_layer']}.",
        font=body_font,
        size=16,
        color=brand,
        bold=True,
        after=16,
    )
    _add_text(
        doc,
        "The site already has a clear point of view. The next gain comes from making the choice, proof, and next step easier.",
        font=body_font,
        size=12,
        color=muted,
    )

    _page_break(doc)
    _add_page_title(doc, "What already works", "Keep the parts that help people understand the business.", display_font, body_font, ink, accent)
    for item in audit["client_review"]["strengths"]:
        _add_bullet(doc, item, body_font, ink)
    _add_text(doc, "What a visitor may think", font=display_font, size=18, color=ink, before=12, after=8)
    for item in audit["client_review"]["visitor_view"]:
        _add_bullet(doc, f"{item['title']}: {item['text']}", body_font, ink)

    _page_break(doc)
    _add_page_title(doc, "Copy review", "The message recognizes the problem. It needs more proof and fewer abstract words.", display_font, body_font, ink, accent)
    if audit["client_review"]["copy"]["quote"]:
        _add_text(
            doc,
            f"“{audit['client_review']['copy']['quote']}”",
            font=display_font,
            size=21,
            color=brand,
            italic=True,
            after=16,
        )
    _add_text(doc, "What works", font=body_font, size=11, color=accent, bold=True)
    for item in audit["client_review"]["copy"]["works"]:
        _add_bullet(doc, item, body_font, ink)
    _add_text(doc, "What weakens it", font=body_font, size=11, color=accent, bold=True, before=8)
    for item in audit["client_review"]["copy"]["risks"]:
        _add_bullet(doc, item, body_font, ink)

    _page_break(doc)
    _add_page_title(doc, "What would make this stronger", "Every score comes with the next move.", display_font, body_font, ink, accent)
    for item in audit["client_review"]["conversion_path"]:
        _add_text(
            doc,
            f"{item['plain_name']} — {item['score']}/{item['max']}",
            font=body_font,
            size=11.5,
            color=brand,
            bold=True,
            after=3,
        )
        _add_text(doc, item["to_10"], font=body_font, size=10.5, color=muted, after=8)

    if audit["mode"] in {"visibility", "full"}:
        _page_break(doc)
        _add_page_title(doc, "Search and AI visibility", "Being easy to quote is not the same as being easy to understand.", display_font, body_font, ink, accent)
        _add_text(doc, audit["client_review"]["visibility_explanation"], font=body_font, size=11, color=muted, after=14)
        for item in audit["client_review"]["visibility_path"]:
            _add_text(
                doc,
                f"{item['name']} — {item['score']}/{item['max']}",
                font=body_font,
                size=11,
                color=brand,
                bold=True,
                after=2,
            )
            _add_text(doc, item["to_10"], font=body_font, size=10, color=muted, after=6)

    _page_break(doc)
    _add_page_title(doc, "First week", "Start with the changes that make the buying decision easier.", display_font, body_font, ink, accent)
    for index, item in enumerate(audit["seven_day_plan"][:5], start=1):
        _add_text(doc, f"{index}. {item}", font=body_font, size=12, color=ink, after=10)
    _add_text(
        doc,
        "This file is editable. Change the wording, add notes, or import it into Google Docs before creating the final client-ready version.",
        font=body_font,
        size=10,
        color=muted,
        italic=True,
        before=16,
    )

    doc.core_properties.title = f"{audit['brand']['name']} Website Review"
    doc.core_properties.subject = "Editable website audit brief"
    doc.core_properties.author = "WebsiteAudit"
    doc.core_properties.keywords = "website audit, conversion, visibility"
    doc.save(path)
    return path
